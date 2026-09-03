# -*- coding: utf-8 -*-
"""Generate KMBA architecture Word document."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(__file__).resolve().parent
KB = (ROOT / 'kb-engine.js').read_text(encoding='utf-8')
VERSION = 'V.0732'
DATE = '2026/07/29'
COMMIT = 'ee91f50'
OUT = ROOT.parent / f'KMBA-CLUB-2026-資料庫與知識庫架構-{VERSION.replace(".", "")}.docx'

KNOWLEDGE_ITEMS = []
for block in re.finditer(r"\{\s*\n\s*id: '([^']+)',\s*\n\s*topic: '([^']+)',", KB):
    KNOWLEDGE_ITEMS.append({'id': block.group(1), 'topic': block.group(2)})

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.font.size = Pt(11)


def h1(text):
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


h1('KMBA CLUB 2026 活動小助手')
para(f'文件版本：{VERSION}')
para(f'建立日期：{DATE}')
para('線上網址：https://zxcvaden-hub.github.io/KMBA-FAQ/')
para('GitHub Repo：https://github.com/zxcvaden-hub/KMBA-FAQ')
para(f'最新 commit：{COMMIT}')

h1('一、系統架構總覽')
para('本專案為純前端靜態 FAQ 客服網站，不含傳統 SQL 資料庫。資料分為三層：')
bullets([
    '前端展示層：index.html + kb-engine.js（GitHub Pages 部署）',
    '知識庫層：kb-engine.js 內 KNOWLEDGE 陣列（靜態 FAQ 規則引擎）',
    '客服紀錄層：Google Sheet「KMBA客服紀錄」+ Google Apps Script Web App',
])
para('另含 lucky/ 雙月抽獎工具（隱藏入口 luckydrawsetting / luckydrawlist，不寫入客服紀錄）。')

h1('二、Google Sheet 客服紀錄（後台資料表）')
para('試算表名稱：KMBA客服紀錄')
para('Spreadsheet ID：1vnxyG4AFXVnfeWkIHVcBctTWWprjMXXmJEFEhEzgcN4')
para('工作表：工作表1（若不存在則使用第一個工作表）')

h2('2.1 欄位結構（A–J）')
table = doc.add_table(rows=11, cols=3)
table.style = 'Table Grid'
headers = ['欄位', '名稱', '說明']
rows = [
    ('A', '時間', '寫入當下時間（Apps Script new Date()）'),
    ('B', 'Session ID', '同一分頁工作階段 ID，格式 KMBA-{timestamp}-{random}'),
    ('C', '使用者問題', '使用者原始提問'),
    ('D', 'AI 回答摘要', '先結論文字，最多 500 字'),
    ('E', '命中 FAQ', 'faqId，如 raffle_fairness、rewards_guide'),
    ('F', 'FAQ 類別', 'category / topic，如 raffleTicket、giftCard'),
    ('G', '是否已解決', '初次空白；回饋「有，已了解」→ YES；「還有其他問題」→ NO'),
    ('H', '裝置', '如 Windows / Chrome、iPhone / Safari'),
    ('I', '版本', f'前端版本，目前 {VERSION}'),
    ('J', '備註', 'sessionId=...;messageId=KMBA-MSG-...;intent=...'),
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row in enumerate(rows, start=1):
    for c, val in enumerate(row):
        table.rows[r].cells[c].text = val

h2('2.2 Web App API')
para('Web App URL（/exec）：')
para('https://script.google.com/macros/s/AKfycbzjTNHPh1mtr7gANezaRj4WC5gTUu-Dm8KzXRtj0ObrlhmxTo_LN_bEnWVUSvwa1ND0GQ/exec')

h2('2.3 Create / Update Payload')
para('Create：{ action:"create", timestamp, sessionId, messageId, question, answer, faqId, category, resolved:"", device, version, note, row[] }')
para('Update：{ action:"update", messageId, resolved:"YES"|"NO", timestamp, sessionId, note }')

h2('2.4 不寫入紀錄的項目')
bullets([
    'luckydrawsetting / luckydrawlist 抽獎隱藏指令',
    '回饋後續訊息（intent: feedback）',
    '歡迎畫面載入（尚未提問）',
])

h1('三、前端 Session / Message ID')
bullets([
    'Session ID：sessionStorage key = kmbaSessionId，頁面載入時建立，同一分頁共用',
    'Message ID：每次提問建立 KMBA-MSG-{timestamp}-{random}，一題一筆',
    '前端常數 KMBA_LOG_ENDPOINT 指向上述 Web App',
    'fetch POST + mode:no-cors，失敗不影響 FAQ 回答',
])

h1('四、FAQ 知識庫（kb-engine.js KNOWLEDGE）')
para(f'共 {len(KNOWLEDGE_ITEMS)} 個知識項目，依 id / topic 如下：')

desc_map = {
    'tasks_guide': '任務類型解說',
    'rewards_guide': '獎勵有哪些（直接完整回答）',
    'gift_card_earn': '商品卡怎麼拿',
    'gift_card_delivery': '商品卡發放／領取方式',
    'gift_card_rank_after_15': '15 名以後商品卡面額',
    'raffle_prizes': '雙月抽獎獎項',
    'raffle_fairness': '抽獎公平／機制／螢幕錄影（V0732 客戶確認文案）',
    'visit_partner_stores': '合作店家／會通知嗎',
    'raffle_100': '100 分抽獎券',
    'raffle_200': '200 分抽獎券',
    'raffle_300': '300 分抽獎券',
    'raffle_rules_full': '完整抽獎規則',
    'raffle_reset': '抽獎券重新計算',
    'visit_points': '拜訪任務積分',
    'visit_one_ticket': '拜訪一間抽獎券',
    'visit_max': '每月最多拜訪',
    'visit_task': '拜訪任務說明',
    'passport_alias': 'Passport 用語導向',
    'customer_photo_bare': '客人推薦照片（短詞）',
    'display_photo_bare': '新品陳列照片（短詞）',
    'visit_photo_bare': '拜訪照片（短詞）',
    'customer_photo': '客人推薦照片規則',
    'display_photo': '新品陳列照片規則',
    'photo_need_new_product': '照片一定要有新品',
    'customer_photo_definition': '客人推薦照片定義',
    'customer_no_face': '客人沒露臉',
    'product_only': '只有產品',
    'gift_card_tiers': '商品卡級距',
    'ranking_tie': '積分相同排名',
    'ranking_calc': '排名計算',
    'visit_has_ticket': '拜訪也有抽獎券',
    'quiz_task': '品牌隨堂考',
    'customer_photo_fail': '照片可能不通過',
    'display_clear': '新品要拍多清楚',
    'display_env': '店內環境',
    'visit_photo_what': '拜訪照片要拍什麼',
    'monthly_tasks': '每月有哪些任務',
}

kt = doc.add_table(rows=len(KNOWLEDGE_ITEMS) + 1, cols=3)
kt.style = 'Table Grid'
kt.rows[0].cells[0].text = 'faqId'
kt.rows[0].cells[1].text = 'topic（類別）'
kt.rows[0].cells[2].text = '說明'
for i, item in enumerate(KNOWLEDGE_ITEMS, start=1):
    kt.rows[i].cells[0].text = item['id']
    kt.rows[i].cells[1].text = item['topic']
    kt.rows[i].cells[2].text = desc_map.get(item['id'], '')

h1('五、追問主題（TOPIC_CLARIFICATIONS）')
bullets([
    'giftCard：商品卡 / 禮券相關追問',
    'raffle：抽獎 / 抽獎券追問',
    'mission：任務類型追問',
    'photo：照片類型追問',
    'visit：拜訪任務追問',
    'points：積分追問',
    'reward：獎勵追問',
    'ranking：排名追問',
])
para('「獎勵有哪些」：直接完整回答，不再走追問（V0729 起）。')

h1('六、核心業務規則摘要')
bullets([
    '對外：各任務100分、月上限300、抽獎券100/200/300；內部排行榜算法見 INTERNAL_RANKING（不公開）',
    '商品卡依每月積分排行：500/200/100 元（前 20/21-40/41-100 名）；101 名起不發放',
    '雙月抽獎：10 月 15×1000 元；12 月 15×2000 元（2026/10/01 前不顯示 12 月場）',
    '拜訪任務：無積分，每店 1 張抽獎券，每月最多 5 間；合作店家請洽區域業務',
    '抽獎公平：台灣彩券種子 + 固定公式 + 10/12 月場全程螢幕錄影並公布',
    '禁止顯示用語：Google、SurveyCake、TOP、Passport、API、資料庫等',
])

h1('七、抽獎工具（lucky/）')
bullets([
    '操作端：chatbot 輸入 luckydrawsetting',
    '觀看端：chatbot 輸入 luckydrawlist（?view=1）',
    '上傳 CSV 名單（姓名、抽獎券數）',
    '公開種子：建議台灣彩券指定期開獎號碼',
    'Efraimidis–Spirakis 加權抽樣（不放回）；同一人最多中獎一次',
    '名單 hash + 種子 → 結果可重現驗證',
])

h1('八、專案檔案結構')
bullets([
    'index.html — 聊天 UI、Session/Message ID、create/update 紀錄、chips UX',
    'kb-engine.js — FAQ 引擎（KNOWLEDGE、追問、多輪對話、照片初步判斷）',
    'google-apps-script-Code.gs — Sheet 寫入後端（手動部署）',
    'kmba-logo.png — LOGO',
    'lucky/index.html、lucky/lucky.js、lucky/banner.png — 抽獎工具',
    'KMBA_CLUB_2026_MASTER_KNOWLEDGE_BASE_V1.0.md — 內部知識庫參考',
    'README.md — 專案說明',
])

h1('九、版本更新紀錄')
bullets([
    f'{VERSION}（{DATE}）：抽獎公平 FAQ 文案優化；本地清理殘留備份',
    'V0731（2026/07/29）：抽獎公平 FAQ（彩券種子、螢幕錄影、網頁公布）',
    'V0730（2026/07/29）：客服紀錄分析 FAQ 缺口修補',
    'V0729（2026/07/29）：UX 滾動/chips；獎勵直接回答；新 FAQ',
    'V0728（2026/07/28）：Google Sheet 客服紀錄',
    'V0724（2026/07/24）：十項 FAQ 改善、抽獎頁改版',
])

h1('十、備份與部署')
bullets([
    f'GitHub Pages commit {COMMIT}（{DATE}）',
    f'架構文件：KMBA-CLUB-2026-資料庫與知識庫架構-{VERSION.replace(".", "")}.docx',
    '交付壓縮檔：大韓菸草客服20260729-還原點.zip',
])

doc.save(OUT)
print(str(OUT))
